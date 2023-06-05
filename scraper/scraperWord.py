import requests
from lxml import html
from docx import Document

def scrape_page(url):
    content = requests.get(url).content
    return process(content)

def process(html_string):
    html_tree = html.fromstring(html_string)
    listings = html_tree.xpath('//div[contains(@class, "vehicle-details")]')
    
    # Create a new Word document
    doc = Document()
    
    for listing in listings:
        mileage_div = listing.xpath('.//div[contains(@class, "mileage")]')
        link = listing.xpath('.//a[@href]')
        price_span = listing.xpath('.//span[contains(@class, "primary-price")]')
        
        mileage = mileage_div[0].text.strip() if mileage_div else 'N/A'
        listing_link = 'https://www.cars.com' + link[0].get('href') if link else 'N/A'
        price = price_span[0].text.strip() if price_span else 'N/A'

        
        # Add the listing details to the document
        doc.add_paragraph("Listing Found:")
        doc.add_paragraph("Mileage: " + mileage)
        doc.add_paragraph("Link: " + listing_link)
        doc.add_paragraph("Price: " + price)

        doc.add_paragraph('-----')

    # Save the document
    doc.save('scraped_results.docx')
    print("Scraped results saved as 'scraped_results.docx'")

url = 'https://www.cars.com/shopping/results/?deal_ratings[]=great&dealer_id=&keyword=&list_price_max=15000&list_price_min=&makes[]=mazda&maximum_distance=100&mileage_max=&models[]=mazda-cx_5&monthly_payment=271&page_size=20&sort=best_deal&stock_type=used&year_max=&year_min=&zip=30064'

scrape_page(url)